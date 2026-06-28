import os
import fitz
from docx import Document as DocxDocument
from langchain_core.documents import Document

class DataToDoc:
    def __init__(self, folder_path="./InputData"):
        self.folder_path = folder_path
        self.documents = []

        pdf_count = 0
        wordfile_count = 0

        for filename in os.listdir(self.folder_path):
            file_path = os.path.join(self.folder_path, filename)

            if filename.lower().endswith(".pdf"):
                pdf = fitz.open(file_path)

                # one document per page
                for page_num, page in enumerate(pdf):
                    self.documents.append(
                        Document(
                            page_content=page.get_text(),
                            metadata={
                                "source": filename,
                                "page": page_num + 1,
                                "type": "pdf"
                            }
                        )
                    )
                pdf.close()
                pdf_count += 1
            elif filename.lower().endswith(".docx"):
                doc = DocxDocument(file_path)
                text = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())

                self.documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": filename,
                            "page": 1,
                            "type": "docx"
                        }
                    )
                )
                wordfile_count += 1

        print("Total PDFs      :", pdf_count)
        print("Total Word Files:", wordfile_count)
        print("Total Documents :", len(self.documents))

    def get_documents(self):
        return self.documents
